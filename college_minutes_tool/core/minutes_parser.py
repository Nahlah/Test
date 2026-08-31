"""
تحليل محاضر المجالس (مجلس القسم أو مجلس الكلية) بالاعتماد على بنية الجداول
الموحّدة المستخدمة في النموذج:
  - جدول الترويسة: اسم المجلس / رقم الجلسة | اليوم | التاريخ | المكان | الوقت
  - "أولاً: جدول الأعمال"
  - "ثانياً: تشكيل الأعضاء" (أو "أعضاء التشكيل")
  - جدول لكل موضوع يحوي: الموضوع NN | العنوان
        حيثيات الموضوع | ...
        التوصية / القرار | ...
        المستند | ...
        المرفقات | ...
        الإجراء المطلوب | ...
        تصويت الأعضاء | الاسم | موافق | غير موافق | ملاحظات  (+ صفوف الأعضاء)
        استضافة | الاسم | الصفة (+ صفوف الضيوف)
  - "رابعاً: توقيع الأعضاء"
  - "خامساً: توجيه صاحب الصلاحية"
"""
import re
from dataclasses import dataclass, field
from docx import Document

from .docx_struct import cell_text, norm, row_texts
from .extract import read_text


LABEL_TOPIC = "الموضوع"
LABEL_RATIONALE = norm("حيثيات الموضوع")
LABEL_DECISION = norm("التوصية / القرار")
LABEL_DOCUMENT = norm("المستند")
LABEL_ATTACH = norm("المرفقات")
LABEL_ACTION = norm("الإجراء المطلوب")
LABEL_VOTE = norm("تصويت الأعضاء")
LABEL_GUEST = norm("استضافة")
LABEL_SESSION_NO = norm("رقم الجلسة")
LABEL_MEMBERS_HDR = "الاسم"


@dataclass
class Topic:
    number: str = ""
    title: str = ""
    rationale: str = ""
    decision: str = ""
    document_ref: str = ""
    attachments: str = ""
    action_required: str = ""
    votes: list = field(default_factory=list)   # [{name, agree, disagree, note}]
    guests: list = field(default_factory=list)  # [{name, role}]


@dataclass
class SessionInfo:
    council_name: str = ""
    session_number: str = ""
    day: str = ""
    date: str = ""
    place: str = ""
    time: str = ""


@dataclass
class Minutes:
    session: SessionInfo = field(default_factory=SessionInfo)
    members: list = field(default_factory=list)  # [{name, rank, role, attendance}]
    topics: list = field(default_factory=list)   # [Topic]
    source_path: str = ""


def _parse_session_header(doc) -> SessionInfo:
    info = SessionInfo()
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            texts = row_texts(row)
            if not texts:
                continue
            if norm(texts[0]) == LABEL_SESSION_NO and len(texts) >= 5:
                # الصف التالي يحوي القيم الفعلية
                if i + 1 < len(table.rows):
                    values = row_texts(table.rows[i + 1])
                    if len(values) >= 5 and any(v.strip() for v in values):
                        info.session_number = values[0].strip()
                        info.day = values[1].strip()
                        info.date = values[2].strip()
                        info.place = values[3].strip()
                        info.time = values[4].strip()
        # اسم المجلس عادة أول صف نصي في أول جدول
    if doc.tables:
        first_row_text = cell_text(doc.tables[0].rows[0].cells[0]) if doc.tables[0].rows else ""
        info.council_name = _clean_council_name(first_row_text)
    return info


def _clean_council_name(text: str) -> str:
    text = text.strip()
    for prefix in ("اسم اللجنة", "اللجنة"):
        if text.startswith(prefix):
            text = text[len(prefix):].strip()
    return text.lstrip("/ ").strip()


def _iter_all_tables(tables):
    """يمرّ على الجداول مع الجداول المتداخلة داخل الخلايا (شائعة في محاضر طويلة الأعضاء)."""
    for table in tables:
        yield table
        for row in table.rows:
            for cell in row.cells:
                if cell.tables:
                    yield from _iter_all_tables(cell.tables)


def _looks_like_roster_header(texts) -> bool:
    joined = "".join(norm(t) for t in texts)
    return norm("الاسم") in joined and norm("الحضور") in joined


def _parse_members(doc) -> list:
    members = []
    seen_names = set()
    for table in _iter_all_tables(doc.tables):
        in_roster = False
        for row in table.rows:
            vals = row_texts(row)
            if not vals:
                continue
            if _looks_like_roster_header(vals):
                in_roster = True
                continue
            if not in_roster:
                continue
            if len(vals) < 2 or not vals[1].strip():
                in_roster = False  # صف فارغ ينهي جدول الأعضاء
                continue
            if norm(vals[0]).startswith(LABEL_TOPIC) or norm(vals[0]) in (LABEL_VOTE, LABEL_GUEST):
                in_roster = False
                continue
            name = vals[1].strip()
            if name in seen_names:
                continue
            seen_names.add(name)
            members.append({
                "seq": vals[0].strip() if len(vals) > 0 else "",
                "name": name,
                "rank": vals[2].strip() if len(vals) > 2 else "",
                "role": vals[3].strip() if len(vals) > 3 else "",
                "attendance": vals[4].strip() if len(vals) > 4 else "",
            })
    return members


def _parse_topic_table(table) -> Topic:
    topic = Topic()
    rows = table.rows
    idx = 0
    n = len(rows)
    while idx < n:
        texts = row_texts(rows[idx])
        if not texts:
            idx += 1
            continue
        label = norm(texts[0])
        if label.startswith(LABEL_TOPIC) and not topic.title:
            m = texts[0].strip()
            topic.number = m.replace("الموضوع", "").strip()
            topic.title = texts[1].strip() if len(texts) > 1 else ""
        elif label == LABEL_RATIONALE:
            topic.rationale = texts[1].strip() if len(texts) > 1 else ""
        elif label == LABEL_DECISION:
            topic.decision = texts[1].strip() if len(texts) > 1 else ""
        elif label == LABEL_DOCUMENT:
            topic.document_ref = texts[1].strip() if len(texts) > 1 else ""
        elif label == LABEL_ATTACH:
            topic.attachments = texts[1].strip() if len(texts) > 1 else ""
        elif label == LABEL_ACTION:
            topic.action_required = texts[1].strip() if len(texts) > 1 else ""
        elif label == LABEL_VOTE:
            idx += 1
            while idx < n:
                vtexts = row_texts(rows[idx])
                if not vtexts or norm(vtexts[0]) == LABEL_GUEST:
                    break
                if len(vtexts) >= 2 and vtexts[1].strip():
                    topic.votes.append({
                        "name": vtexts[1].strip(),
                        "agree": vtexts[2].strip() if len(vtexts) > 2 else "",
                        "disagree": vtexts[3].strip() if len(vtexts) > 3 else "",
                        "note": vtexts[4].strip() if len(vtexts) > 4 else "",
                    })
                idx += 1
            continue
        elif label == LABEL_GUEST:
            idx += 1
            while idx < n:
                gtexts = row_texts(rows[idx])
                if not gtexts:
                    break
                if len(gtexts) >= 2 and gtexts[1].strip() and norm(gtexts[1]) not in ("لايوجد", "لا يوجد".replace(" ", "")):
                    topic.guests.append({"name": gtexts[1].strip(), "role": gtexts[2].strip() if len(gtexts) > 2 else ""})
                idx += 1
            continue
        idx += 1
    return topic


def _is_topic_table(table) -> bool:
    if not table.rows:
        return False
    for row in table.rows[:2]:
        texts = row_texts(row)
        if texts and norm(texts[0]).startswith(LABEL_TOPIC):
            return True
    return False


def parse_minutes(path: str) -> Minutes:
    """يحلّل ملف Word .docx يتّبع بنية الجداول الموحّدة (الأسلوب الدقيق والموصى به)."""
    doc = Document(path)
    minutes = Minutes(source_path=path)
    minutes.session = _parse_session_header(doc)
    minutes.members = _parse_members(doc)
    for table in doc.tables:
        if _is_topic_table(table):
            topic = _parse_topic_table(table)
            if topic.title:
                minutes.topics.append(topic)
    return minutes


_TOPIC_SPLIT_RE = re.compile(
    r"(?:^|\n)\s*(?:الموضوع|البند)\s*[:\-]?\s*"
    r"([0-9]{1,2}|الأول|الثاني|الثالث|الرابع|الخامس|السادس|السابع|الثامن|التاسع|العاشر)\b"
)


def parse_minutes_plain(path: str) -> Minutes:
    """تحليل تقريبي (Best-effort) لملفات PDF أو نص عادي لا تحمل بنية جداول Word.
    يعتمد على تقسيم النص عند ظهور كلمة 'الموضوع'/'البند'، ويضع بقية كل قسم في الحيثيات
    ليراجعها المستخدم أو يستكملها النموذج المحلي قبل التصدير.
    """
    text = read_text(path)
    minutes = Minutes(source_path=path)
    matches = list(_TOPIC_SPLIT_RE.finditer(text))
    for i, m in enumerate(matches):
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        chunk = text[start:end].strip()
        lines = [ln.strip() for ln in chunk.splitlines() if ln.strip()]
        title = lines[0] if lines else ""
        rationale = "\n".join(lines[1:])
        minutes.topics.append(Topic(number=m.group(1), title=title, rationale=rationale))
    return minutes


def parse_minutes_file(path: str) -> Minutes:
    """نقطة الدخول الموصى بها: تختار طريقة التحليل المناسبة حسب صيغة الملف."""
    if path.lower().endswith(".docx"):
        return parse_minutes(path)
    return parse_minutes_plain(path)
