"""بناء محضر مجلس الكلية النهائي عبر تعبئة النموذج (القالب) الأصلي بأقسام المواضيع المولّدة،
مع الحفاظ على تنسيق القالب كما هو (الجداول، الخطوط، اتجاه RTL) دون إعادة إنشائه من الصفر.
"""
from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.table import Table

from .docx_struct import norm, row_texts, set_cell_text
from .similarity import normalize_ar

LABEL_RATIONALE = norm("حيثيات الموضوع")
LABEL_RATIONALE_ALT = norm("حيثيات/الموضوع")
LABEL_DECISION = norm("التوصية / القرار")
LABEL_DOCUMENT = norm("المستند")
LABEL_ATTACH = norm("المرفقات")
LABEL_ACTION = norm("الإجراء المطلوب")
LABEL_SESSION_NO = norm("رقم الجلسة")
LABEL_SIGNATURE = norm("التوقيع")


def _is_topic_block_table(table) -> bool:
    if not table.rows:
        return False
    texts = row_texts(table.rows[0])
    return bool(texts) and norm(texts[0]).startswith("الموضوع")


def _is_signature_table(table) -> bool:
    for row in table.rows[:3]:
        texts = row_texts(row)
        if any(norm(t) == LABEL_SIGNATURE for t in texts):
            return True
    return False


def _find_topic_block_tables(doc):
    return [t for t in doc.tables if _is_topic_block_table(t)]


def _find_signature_table(doc):
    for t in doc.tables:
        if _is_signature_table(t):
            return t
    return None


def ensure_topic_block_count(doc, needed: int):
    """يضبط عدد جداول (كتل) المواضيع في القالب ليطابق العدد المطلوب، بنسخ/حذف كتلة النموذج."""
    blocks = _find_topic_block_tables(doc)
    if not blocks:
        raise ValueError("لم يتم العثور على أي كتلة 'الموضوع' في ملف القالب.")
    sig_table = _find_signature_table(doc)
    if sig_table is None:
        raise ValueError("لم يتم العثور على جدول توقيع الأعضاء في ملف القالب.")

    if len(blocks) > needed:
        for extra in blocks[needed:]:
            extra._tbl.getparent().remove(extra._tbl)
        blocks = blocks[:needed]
    elif len(blocks) < needed:
        template_tbl_el = blocks[-1]._tbl
        anchor = sig_table._tbl
        for _ in range(needed - len(blocks)):
            new_tbl_el = deepcopy(template_tbl_el)
            anchor.addprevious(new_tbl_el)
            spacer = OxmlElement("w:p")
            new_tbl_el.addnext(spacer)
            blocks.append(Table(new_tbl_el, doc))
    return blocks


def fill_topic_block(table, number_label: str, title: str, rationale: str, decision: str,
                      document_ref: str, attachments: str, action_required: str):
    rows = table.rows
    if not rows:
        return
    set_cell_text(rows[0].cells[0], f"الموضوع {number_label}")
    if len(rows[0].cells) > 1:
        set_cell_text(rows[0].cells[1], title)
    field_map = {
        LABEL_RATIONALE: rationale,
        LABEL_RATIONALE_ALT: rationale,
        LABEL_DECISION: decision,
        LABEL_DOCUMENT: document_ref,
        LABEL_ATTACH: attachments,
        LABEL_ACTION: action_required,
    }
    for row in rows[1:]:
        texts = row_texts(row)
        if not texts:
            continue
        label = norm(texts[0])
        if label in field_map and len(row.cells) > 1:
            set_cell_text(row.cells[1], field_map[label])


def fill_header(doc, session_number: str, day: str, date: str, place: str, time: str):
    values = [session_number, day, date, place, time]
    for table in doc.tables:
        rows = table.rows
        for i, row in enumerate(rows):
            texts = row_texts(row)
            if texts and norm(texts[0]) == LABEL_SESSION_NO and i + 1 < len(rows):
                target = rows[i + 1]
                for cell, val in zip(target.cells, values):
                    set_cell_text(cell, val)


def fill_agenda(doc, topic_titles: list):
    for table in doc.tables:
        rows = table.rows
        for i, row in enumerate(rows):
            texts = row_texts(row)
            if texts and "جدولالاعمال" in normalize_ar(texts[0]).replace(" ", "") and i + 1 < len(rows):
                content_row = rows[i + 1]
                if len(content_row.cells) > 1:
                    set_cell_text(content_row.cells[1], "\n".join(topic_titles))
                elif content_row.cells:
                    set_cell_text(content_row.cells[0], "\n".join(topic_titles))
                return


def build_minutes_document(template_path: str, output_path: str, session_meta: dict,
                            topics: list, start_number: int = 1):
    """
    session_meta: {"session_number","day","date","place","time"}
    topics: قائمة قواميس بالمفاتيح: title, rationale, decision, document_ref, attachments, action_required
    """
    doc = Document(template_path)
    fill_header(doc, session_meta.get("session_number", ""), session_meta.get("day", ""),
                session_meta.get("date", ""), session_meta.get("place", ""), session_meta.get("time", ""))
    fill_agenda(doc, [t["title"] for t in topics])
    blocks = ensure_topic_block_count(doc, len(topics))
    for i, (block, topic) in enumerate(zip(blocks, topics)):
        number_label = f"{start_number + i:02d}"
        fill_topic_block(
            block, number_label,
            title=topic.get("title", ""),
            rationale=topic.get("rationale", ""),
            decision=topic.get("decision", ""),
            document_ref=topic.get("document_ref", ""),
            attachments=topic.get("attachments", ""),
            action_required=topic.get("action_required", ""),
        )
    doc.save(output_path)
