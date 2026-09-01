"""أدوات مساعدة للتعامل مع بنية ملفات Word (فقرات وجداول) بترتيبها الفعلي في المستند."""
import re
import unicodedata
from copy import deepcopy

from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


def iter_block_items(parent):
    """يمرّ على الفقرات والجداول بالترتيب الذي تظهر به فعليًا في المستند."""
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("parent يجب أن يكون Document أو Cell")

    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def norm(text: str) -> str:
    """يزيل المسافات والأسطر الفارغة والشرطة المائلة لمقارنة تسميات الحقول، لأن بعض
    المحاضر تكتب نفس التسمية بصيغ مختلفة (مثل 'حيثيات الموضوع' أو 'حيثيات\nالموضوع'
    أو 'حيثيات/الموضوع'). إغفال هذا التطبيع يجعل الحقل يبدو فارغًا تمامًا رغم وجود
    محتواه فعليًا في الجدول، بمجرد اختلاف الفاصل المستخدم بين كلمتي التسمية."""
    if text is None:
        return ""
    text = re.sub(r"\s+", "", text)
    return text.replace("/", "")


def cell_text(cell: _Cell) -> str:
    """نص كل فقرات الخلية مفصولة بسطر جديد (يحافظ على الفواصل بين الأسطر داخل نفس الخلية).
    يُطبَّع النص عبر NFKC لتصحيح بعض النصوص العربية التي تُخزَّن أحيانًا بأشكال عرض حروف
    منفصلة (Presentation Forms) داخل مربعات نص/WordArt فتظهر بلا مسافات صحيحة.
    """
    text = "\n".join(p.text for p in cell.paragraphs).strip()
    return unicodedata.normalize("NFKC", text)


def row_texts(row) -> list:
    return [cell_text(c) for c in row.cells]


def clear_cell(cell: _Cell):
    """يفرغ محتوى الخلية ويبقي فقرة واحدة فارغة بنفس التنسيق الافتراضي للفقرة الأولى."""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p0 = cell.paragraphs[0]
    for run in list(p0.runs):
        run._element.getparent().remove(run._element)


def set_cell_text(cell: _Cell, text: str):
    """يستبدل محتوى الخلية بنص جديد (قد يشمل أسطرًا متعددة مفصولة بـ \n)، مع الحفاظ على
    تنسيق الفقرة الأصلية (بما في ذلك اتجاه الكتابة من اليمين لليسار/المحاذاة) في كل الأسطر.

    فقرات Word الجديدة التي تُنشأ عبر add_paragraph() لا ترث خصائص الفقرة المباشرة (كاتجاه
    RTL أو المحاذاة) الموجودة على الفقرة الأصلية في القالب، فتظهر معكوسة الاتجاه (يسار
    لليمين) بصريًا رغم أن النص عربي. الحل: نسخ عنصر <w:pPr> نفسه من الفقرة الأولى لكل فقرة
    إضافية بدل الاعتماد على الأسلوب (style) فقط.
    """
    lines = (text or "").split("\n")
    clear_cell(cell)
    p0 = cell.paragraphs[0]
    base_font = None
    try:
        base_font = cell.paragraphs[0].style
    except Exception:
        base_font = None
    p0_pPr = p0._p.find(qn("w:pPr"))

    run = p0.add_run(lines[0])
    for extra in lines[1:]:
        p = cell.add_paragraph()
        if base_font is not None:
            p.style = base_font
        if p0_pPr is not None:
            existing_pPr = p._p.find(qn("w:pPr"))
            if existing_pPr is not None:
                p._p.remove(existing_pPr)
            p._p.insert(0, deepcopy(p0_pPr))
        p.add_run(extra)
