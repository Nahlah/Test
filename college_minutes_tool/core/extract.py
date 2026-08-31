"""استخراج النص من ملفات Word / PDF / نص عادي."""
import os

from docx import Document
import pdfplumber


def read_text(path: str) -> str:
    """يقرأ محتوى الملف كنص عادي بغض النظر عن صيغته (docx / pdf / txt)."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return _read_docx(path)
    if ext == ".pdf":
        return _read_pdf(path)
    return _read_plain(path)


def _read_docx(path: str) -> str:
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    # يشمل أيضًا أي نص داخل الجداول (شائع في محاضر المجالس)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _read_pdf(path: str) -> str:
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            parts.append(text)
    return "\n".join(parts)


def _read_plain(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()
